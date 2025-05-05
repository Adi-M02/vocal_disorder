import json
from pathlib import Path
import re
import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# === CONFIG ===
INPUT_JSON = Path("vocab_output_04_20/expanded_vocab_ngram_top40_uf5_bf3_tf2.json")

def format_vocab_for_word(vocab_dict):
    lines = []

    lines.append("Expanded Vocabulary by Category")
    lines.append("=" * 35)
    lines.append("")

    for category, terms in sorted(vocab_dict.items()):
        lines.append(f"{category} ({len(terms)} terms)")
        lines.append("-" * len(lines[-1]))

        sorted_terms = sorted(set(terms), key=lambda x: x.lower())
        for term in sorted_terms:
            readable_term = term.replace("_", " ")  # Replace underscores with spaces
            lines.append(readable_term)
        lines.append("")  # add spacing between sections

    return "\n".join(lines)

def highlight_terms(file_path, terms):
    """
    Reads the UTF-8 text file at `file_path`, highlights all whole-word matches
    of each term in `terms` (case-insensitive), and writes a .docx at
    file_path.replace('.txt', '_highlighted.docx').
    """
    # 1. Load lines
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 2. Build regex: sort by length so longer terms win over substrings
    terms = sorted(terms, key=len, reverse=True)
    pattern = re.compile(r'\b(' + '|'.join(re.escape(t) for t in terms) + r')\b',
                         flags=re.IGNORECASE)

    # 3. Create Word doc and walk each line
    doc = Document()
    for line in lines:
        p = doc.add_paragraph()
        last = 0
        for m in pattern.finditer(line):
            # plain text before the term
            p.add_run(line[last:m.start()])
            # highlighted term
            run = p.add_run(m.group(0))
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            last = m.end()
        # any leftover text after last match (including newline)
        p.add_run(line[last:])

    # 4. Save to <original>_highlighted.docx
    base, _ = os.path.splitext(file_path)
    out_path = f"{base}_highlighted.docx"
    doc.save(out_path)
    print(f"Highlighted doc saved to: {out_path}")
    return out_path

def main():
    with INPUT_JSON.open("r", encoding="utf-8") as f:
        data = json.load(f)

    vocab_dict = data.get("vocabulary", {})

    formatted_output = format_vocab_for_word(vocab_dict)

    # Create output path in the same directory with modified name
    output_path = INPUT_JSON.with_name(INPUT_JSON.stem + "_word_friendly.txt")

    with output_path.open("w", encoding="utf-8") as f:
        f.write(formatted_output)

    print(f"✅ Saved Word-friendly output to: {output_path}")

if __name__ == "__main__":
    # main()
    with open('vocabulary_evaluation/manual_terms.txt', 'r', encoding='utf-8') as f:
        content = f.read()
        terms = [term.strip() for term in content.split(',') if term.strip()]
    highlight_terms('vocabulary_evaluation/all_ground_truth_users.txt', terms)